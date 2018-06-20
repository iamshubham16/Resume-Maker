#PROJECT#
#CV BUILDER#
#BY Shubham Guglani (2016KUCP1018)#
#Shrishti Gupta  (2016KUCP1019)#
#Branch - CSE#


from Tkinter import *                     #library to create GUI
from tkFileDialog import askopenfilename  #to import open file dialog box
from docx import Document                 #library to import document
from docx.shared import Inches            #to resize the images stored in doc file
import sys
import os
import comtypes.client                    #to convert .doc to .pdf
from PIL import ImageTk, Image            #to manipulate images.


document = Document()                        #creates new empty document
document.add_heading('Curriculum Vitae',0)   #displays a heading in the doc. file


#field names stored in a tuple
personal_info =  'Name', 'Date of Birth','Nationality', 'Contact', 'Email'  
qualifications = 'School','Duration','University','Duration','Degree'
skills = 'Interests','Comp. Languages'
exp = 'Company1','Post','Duration','Company2','Post','Duration'

#function called when 'Preview' button is clicked.
#it displays the "uploaded image" in a popup window
def display():
    popup = Toplevel()  #creates the toplevel for popup window
    popup.wm_title("Image") #assigns title to the popup window
    path = image            #assigns the path of the image
    img = ImageTk.PhotoImage(Image.open(path)) 
    w1 = Label(popup,image=img).pack(side=LEFT) #packs the image on the toplevel window
    popup.mainloop()

#function called when 'Upload Image' button is clicked.
#it opens the 'Open File' dialog box.
def callback():
   global image 
   image=askopenfilename()    #opens the dialog box to select file from any directory.
   print "Image Uploaded"
   btn = btn_text.get()       #gets the current value of the button
   if btn=='Upload Image':    #if current value == upload image, then preview button is created 
       b3 = Button(root, text='Preview',command=display) #button when clicked, calls display() function.
       b3.pack(side=LEFT, padx=6, pady=6)
   btn_text.set("Change Image") #once the image is uploaded,changes the button text to 'Change Image'.
   

#function called when create CV button is clicked.
#function creates the .doc/.pdf file with the details entered.
def fetch(entries,root,i):
    l=[]                    #creates an empty list 
    for entry in entries:    #iterates over the list of values received from the form
      field = entry[0]
      text  = entry[1].get()
      l.append(text)          #appends the values of each field in a list.
    root.destroy()  
    document.add_picture(image,width=Inches(2.25))          #inserts the image in the doc file.
    document.add_paragraph("Personal Information: ",style="Intense Quote")  #inserts user's personal details in the doc file
    document.add_paragraph("Name :   ").add_run(l[0])
    document.add_paragraph("D.O.B :   ").add_run(l[1])
    document.add_paragraph("Nationality :   ").add_run(l[2])
    document.add_paragraph("Objective : ",style="Intense Quote") #displays user's objective
    document.add_paragraph("To acquire valuable knowledge and skills to complement that I have learnt from school in an actual job environment.")
    document.add_paragraph("In return, I offer my service and determination to be an asset to your company throughout the duration of my training period.")
    document.add_paragraph('Educational Qualifications: ',style="Intense Quote")  #inserts user's educational info. in the doc file
    document.add_paragraph("School :   ").add_run(l[5])
    document.add_paragraph("Duration :   ").add_run(l[6])
    document.add_paragraph().add_run("Degree :   "+l[9])
    document.add_paragraph("College/University :   ").add_run(l[7])
    document.add_paragraph("Duration :   ").add_run(l[8])
    document.add_paragraph('Skills: ',style="Intense Quote")            #inserts user's skills and interests in the doc file
    document.add_paragraph("Interests :   ",style='List Bullet').add_run().bold=True
    document.add_paragraph("--> "+l[10])
    document.add_paragraph("Computer Languages Known :   ",style='List Bullet').add_run().bold=True
    document.add_paragraph("--> "+l[11])
    document.add_paragraph("Work Experience: ",style="Intense Quote")   #inserts user's work experience
    document.add_paragraph('1. ',style='List Bullet').add_run(l[12]+' ( '+l[13]+' )').bold=True
    document.add_paragraph("Duration :   ").add_run(l[14])
    document.add_paragraph('2. ',style='List Bullet').add_run(l[15]+' ( '+l[16]+' )').bold=True
    document.add_paragraph("Duration :   ").add_run(l[17])

    document.add_paragraph("Contact Details: ",style="Intense Quote")   #displays user's contact details in the doc file
    document.add_paragraph("Mobile:  ").add_run(l[3])
    document.add_paragraph("Email:  ").add_run(l[4])
    if i==1:
       document.save(l[0]+'_CV'+'.docx')   #creates CV in .doc format
    else:
       document.save(l[0]+'_CV'+'.docx')
       wdFormatPDF = 17
       in_file = os.path.abspath(l[0]+'_CV'+'.docx')
       out_file = os.path.abspath(l[0]+'_CV'+'.pdf') #creates CV in .pdf format as per user's choice

       word = comtypes.client.CreateObject('Word.Application')
       doc = word.Documents.Open(in_file)
       doc.SaveAs(out_file, FileFormat=wdFormatPDF) #saves the pdf with the file name consisting of user's name.
       doc.Close()
       word.Quit()

    print "CV Created"
    
def makeform(root,fields,qual,skills,exp):      #function to create GUI
   entries = []
   font1 = "-family {@Times} -size 10 -weight bold -slant roman -underline 1 -overstrike 0"
   L1=Label(root, text="Personal Information",fg="black",font=font1).pack(side =TOP) #creates label displaying 'personal information'
   for field in fields:     #a loop that creates 'entry' widgets dynamically under 'personal information' category.
      row = Frame(root) 
      row.configure(borderwidth="2")
      lab = Label(row, width=15, text=field, anchor='w')
      lab.place(relx=0.02,rely=0.09,height=31,width=73)
      lab.configure(background="#d9d9d9")
      lab.configure(disabledforeground="#a3a3a3")
      lab.configure(foreground="#000000")
      ent = Entry(row)
      ent.place(relx=0.18,rely=0.09,relheight=0.04,relwidth=0.29)
      ent.configure(background="white")
      ent.configure(disabledforeground="#a3a3a3")
      ent.configure(foreground="#000000")
      row.pack(side=TOP, fill=X, padx=5, pady=5)
      lab.pack(side=LEFT)
      ent.pack(side=LEFT,expand=YES,fill=X)
      entries.append((field, ent))       #appends the field and its corresponding value in the list.
   Label(root, text="Educational Qualifications",fg="black",font=font1).pack()  #creates label displaying 'educational info.'
   for q in qual:
      row = Frame(root)         #a loop that creates 'entry' widgets dynamically under 'educational qualifications' category.
      lab = Label(row, width=15, text=q, anchor='w')
      lab.place(relx=0.02,rely=0.54,height=31,width=73)
      lab.configure(background="#d9d9d9")
      lab.configure(disabledforeground="#a3a3a3")
      lab.configure(foreground="#000000")
      ent = Entry(row)
      ent.place(relx=0.18,rely=0.54,relheight=0.04,relwidth=0.29)
      ent.configure(background="white")
      ent.configure(disabledforeground="#a3a3a3")
      ent.configure(foreground="#000000")
      row.pack(side=TOP, fill=X, padx=5, pady=5)
      lab.pack(side=LEFT)
      ent.pack(side=LEFT,expand=YES,fill=X)
      entries.append((q, ent))         #appends the field and its corresponding value in the list.
   Label(root, text="Skills",fg="black",font=font1).pack()      #creates label displaying 'skills'
   for s in skills:                     #a loop that creates 'entry' widgets dynamically under 'skills' category.
      row = Frame(root)
      lab = Label(row, width=15, text=s, anchor='w')
      lab.place(relx=0.02,rely=0.09,height=31,width=73)
      lab.configure(background="#d9d9d9")
      lab.configure(disabledforeground="#a3a3a3")
      lab.configure(foreground="#000000")
      ent = Entry(row)
      ent.place(relx=0.18,rely=0.09,relheight=0.04,relwidth=0.29)
      ent.configure(background="white")
      ent.configure(disabledforeground="#a3a3a3")
      ent.configure(foreground="#000000")
      row.pack(side=TOP, fill=X, padx=5, pady=5)
      lab.pack(side=LEFT)
      ent.pack(side=LEFT,expand=YES,fill=X)
      entries.append((s, ent))              #appends the field and its corresponding value in the list.
   Label(root, text="Work Experience",fg="black",font=font1).pack()  #creates label displaying 'work experience'
   for e in exp:                            #a loop that creates 'entry' widgets dynamically under 'work experience' category.
      row = Frame(root)
      lab = Label(row, width=15, text=e, anchor='w')
      lab.place(relx=0.02,rely=0.09,height=31,width=73)
      lab.configure(background="#d9d9d9")
      lab.configure(disabledforeground="#a3a3a3")
      lab.configure(foreground="#000000")
      ent = Entry(row)
      ent.place(relx=0.18,rely=0.09,relheight=0.04,relwidth=0.29)
      ent.configure(background="white")
      ent.configure(disabledforeground="#a3a3a3")
      ent.configure(foreground="#000000")
      row.pack(side=TOP, fill=X, padx=5, pady=5)
      lab.pack(side=LEFT)
      ent.pack(side=LEFT,expand=YES,fill=X)
      entries.append((e, ent))             #appends the field and its corresponding value in the list.
   return entries                          #returns the values entered in the form.


if __name__ == '__main__':
   root = Tk()              #creates the toplevel window for the gui where widgets are placed
   root.title('CV Generator')  #assigns title to the window
   ents = makeform(root,personal_info,qualifications,skills,exp)  #calls the function to create and display form.
   root.bind('<Return>', (lambda event, e=ents: fetch(e)))        #calls the function to fetch form values and print in the .doc/.pdf file.
   btn_text = StringVar()
   btn_text.set("Upload Image")         #initially sets the button text to 'Upload Image'
   b=Button(textvariable=btn_text,command=callback)
   b.pack(side=LEFT,padx=6, pady=6)
   b1 = Button(root, text='Create CV(.doc)',command=(lambda e=ents: fetch(e,root,1)))   #creates button to generate CV in .doc format
   b1.pack(side=LEFT, padx=6, pady=6)
   b3 = Button(root, text='Create CV(.pdf)',command=(lambda e=ents: fetch(e,root,2)))   #creates button to generate CV in .pdf format
   b3.pack(side=LEFT, padx=6, pady=6)
   b2 = Button(root, text='Quit', command=root.destroy)  #creates Quit button,which destroys the window when clicked.
   b2.pack(side=LEFT, padx=6, pady=6)
   
   root.mainloop()
   
 
#The End :)
